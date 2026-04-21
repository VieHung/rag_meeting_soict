import io
from pypdf import PdfReader
from docx import Document
from pathlib import Path
from typing import Union


class DocumentParser:

    def parse(self, file_bytes: bytes, filename: str, mime_type: str) -> str:
        ext = Path(filename).suffix.lower()

        if ext == ".pdf":
            return self._parse_pdf(file_bytes)
        elif ext in (".docx", ".doc"):
            return self._parse_docx(file_bytes)
        elif ext in (".txt", ".md", ".rst"):
            return self._parse_txt(file_bytes)
        else:
            try:
                return self._parse_txt(file_bytes)
            except Exception:
                raise ValueError(f"Unsupported file type: {ext}")

    def _parse_txt(self, data: bytes) -> str:
        try:
            return data.decode("utf-8")
        except UnicodeDecodeError:
            return data.decode("latin-1", errors="replace")

    def _parse_pdf(self, data: bytes) -> str:
        text_parts = []
        reader = PdfReader(io.BytesIO(data))
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text.strip():
                text_parts.append(f"[Page {page_num + 1}]\n{text}")
        return "\n\n".join(text_parts)

    def _parse_docx(self, data: bytes) -> str:
        doc = Document(io.BytesIO(data))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        table_texts = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    table_texts.append(row_text)

        return "\n\n".join(paragraphs + table_texts)